


from docx import Document

# Question data
questions = [
    {
        "title": "Identifying a Point on a Line Graph",
        "description": "This question assesses understanding of coordinate geometry and interpreting points on a line graph.",
        "question": "The line shown in the coordinate plane passes through the points $(0, 2)$ and $(4, 10)$. Which point lies on the line?",
        "image": "https://i.imgur.com/qZf77FJ.png",
        "instruction": "Select the correct point from the options below.",
        "difficulty": "moderate",
        "order": 1,
        "options": ["(2, 5)", "(3, 7)", "@@(2, 6)", "(1, 4)"],
        "explanation": "The slope of the line is (10-2)/(4-0) = 2. Equation: y = 2x + 2. Substituting x=2 → y=6. So (2,6) lies on the line.",
        "subject": "Quantitative Math",
        "unit": "Geometry and Measurement",
        "topic": "Coordinate Geometry",
        "marks": 1
    },
    {
        "title": "Solving Quadratic Equations",
        "description": "This question tests the ability to solve quadratic equations by factoring.",
        "question": "Solve for $x$: $$x^{2} - 7x + 12 = 0$$",
        "image": None,
        "instruction": "Choose the correct values of $x$.",
        "difficulty": "easy",
        "order": 2,
        "options": ["$x=2$ or $x=3$", "@@$x=3$ or $x=4$", "$x=4$ or $x=5$", "$x=6$ or $x=2$"],
        "explanation": "Factorize: $x^{2} - 7x + 12 = (x-3)(x-4) = 0$. So, $x=3$ or $x=4$.",
        "subject": "Quantitative Math",
        "unit": "Algebra",
        "topic": "Quadratic Equations & Functions (Finding roots/solutions, graphing)",
        "marks": 1
    }
]


doc = Document()

for q in questions:
    doc.add_paragraph(f"@title {q['title']}")
    doc.add_paragraph(f"@description {q['description']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"@question {q['question']}")
    if q['image']:
        doc.add_paragraph(f"![]({q['image']})")
    doc.add_paragraph("")
    doc.add_paragraph(f"@instruction {q['instruction']}")
    doc.add_paragraph(f"@difficulty {q['difficulty']}")
    doc.add_paragraph(f"@Order {q['order']}")
    for opt in q['options']:
        doc.add_paragraph(f"@option {opt}")
    doc.add_paragraph("")
    doc.add_paragraph(f"@explanation {q['explanation']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"@subject {q['subject']}")
    doc.add_paragraph(f"@unit {q['unit']}")
    doc.add_paragraph(f"@topic {q['topic']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"@plusmarks {q['marks']}")
    doc.add_paragraph("")

doc.save("Generated_Math_Questions.docx")


with open("Generated_Math_Questions.md", "w") as f:
    f.write("# Generated Math Questions\n\n")
    for q in questions:
        f.write(f"## Question {q['order']}\n\n")
        f.write(f"@title {q['title']}\n\n")
        f.write(f"@description {q['description']}\n\n")
        f.write(f"@question {q['question']}\n\n")
        if q['image']:
            f.write(f"![]({q['image']})\n\n")
        f.write(f"@instruction {q['instruction']}\n")
        f.write(f"@difficulty {q['difficulty']}\n")
        f.write(f"@Order {q['order']}\n\n")
        for opt in q['options']:
            f.write(f"@option {opt}\n")
        f.write("\n")
        f.write(f"@explanation {q['explanation']}\n\n")
        f.write(f"@subject {q['subject']}\n")
        f.write(f"@unit {q['unit']}\n")
        f.write(f"@topic {q['topic']}\n\n")
        f.write(f"@plusmarks {q['marks']}\n\n---\n\n")

print("Files Generated: Generated_Math_Questions.docx, Generated_Math_Questions.md")
